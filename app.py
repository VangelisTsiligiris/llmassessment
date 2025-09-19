import os, io, json, time, datetime, random, string, html as _html
from collections.abc import Mapping

import streamlit as st
from streamlit_quill import st_quill
from streamlit.components.v1 import html as st_html
import pandas as pd

# ---------- Optional libs ----------
try:
    import gspread
    from gspread.exceptions import WorksheetNotFound
    from google.oauth2.service_account import Credentials
except Exception:
    gspread = None
    WorksheetNotFound = Exception
    Credentials = None

try:
    import google.generativeai as genai
except Exception:
    genai = None

try:
    import docx  # python-docx
    DOCX_OK = True
except Exception:
    DOCX_OK = False

# HTML → text
try:
    from bs4 import BeautifulSoup
    def html_to_text(html: str) -> str:
        return BeautifulSoup(html or "", "html.parser").get_text("\n")
except Exception:
    def html_to_text(html: str) -> str:
        return (html or "").replace("<br>", "\n").replace("<br/>", "\n")

# Markdown renderer (assistant messages)
try:
    import markdown as _md
except Exception:
    _md = None

def md_to_html(text: str) -> str:
    if not text:
        return ""
    if _md:
        try:
            return _md.markdown(
                text,
                extensions=["fenced_code", "tables", "sane_lists", "codehilite"],
            )
        except Exception:
            pass
    # Fallback: escape + bold + newlines
    import re, html as _h
    t = _h.escape(text)
    t = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", t)
    return t.replace("\n", "<br>")

# ---------- Similarity backends ----------
SIM_BACKEND = "none"
_sbert_model = None
try:
    from sentence_transformers import SentenceTransformer, util as sbert_util
    @st.cache_resource
    def load_sbert_model():
        return SentenceTransformer("all-MiniLM-L6-v2")
    _sbert_model = load_sbert_model()
    SIM_BACKEND = "sbert"
except Exception:
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        SIM_BACKEND = "tfidf"
    except Exception:
        from difflib import SequenceMatcher
        SIM_BACKEND = "difflib"

# ---------- Page config & CSS ----------
st.set_page_config(
    page_title="LLM Coursework Helper",
    layout="wide",
    menu_items={"Get help": None, "Report a bug": None, "About": None},
)

st.markdown("""
<style>
:root {
  --ui-font: ui-sans-serif, system-ui, -apple-system, "Segoe UI", Roboto, Helvetica, Arial, "Noto Sans", "Liberation Sans", sans-serif;
  --brand: #3558ff;
  --brand-2: #7c4dff;
  --bg-soft: #f6f8ff;
  --card: #ffffff;
  --muted: #6b7280;
  --border: #e6e9f2;
}

/* Use our UI font (do NOT override icon fonts) */
html, body, [data-testid="stAppViewContainer"] {
  font-family: var(--ui-font);
}

.block-container { padding-top: 0.8rem; padding-bottom: 1rem; }

/* Header chips */
.header-bar {display:flex; gap:.6rem; flex-wrap:wrap; font-size:.95rem; color:#444; margin:.25rem 0 .5rem;}
.status-chip{background:#f5f7fb;border:1px solid var(--border);border-radius:999px;padding:.15rem .6rem}
.small-muted{color:#7a7f8a}

/* Landing */
.hero {
  max-width: 1100px; margin: .6rem auto 1.2rem; padding: 1.25rem 1.4rem;
  border-radius:16px; border:1px solid #dfe6ff;
  background: linear-gradient(135deg, #edf1ff 0%, #f7f4ff 100%);
}
.hero h1 { margin:.1rem 0 .2rem; font-size:2.2rem; line-height:1.2; }
.hero p { color:#334155; margin:.2rem 0 0; }

.info-grid { display:grid; grid-template-columns: 1fr 1fr; gap: 1rem; margin-top: .8rem; }
.card {
  background: var(--card); border:1px solid var(--border); border-radius:14px; padding:1rem 1.1rem;
  box-shadow: 0 1px 0 rgba(17, 24, 39, 0.03);
}
.card h3 { margin:.2rem 0 .6rem; }
.card ul { margin:.25rem 0 .2rem 1.1rem; }
.badge {
  display:inline-block; padding:.2rem .5rem; font-size:.8rem; border-radius:999px;
  background:#eef2ff; color:#334155; border:1px solid #e5e7eb; margin-left:.5rem;
}

/* Chat */
.chat-box { height: 560px; overflow-y:auto; border:1px solid #dcdfe6; border-radius:10px; background:#fff; padding:.5rem; }
.chat-empty{ border:1px dashed #e6e9f2; background:#fbfbfb; color:#708090; padding:.6rem .8rem; border-radius:10px; }
.chat-bubble { border-radius:12px; padding:.7rem .9rem; margin:.45rem .2rem; border:1px solid #eee; line-height:1.55; font-size:0.95rem; }
.chat-user      { background:#eef7ff; }
.chat-assistant { background:#f6f6f6; }
.chat-bubble p { margin:.35rem 0; }
.chat-bubble ul, .chat-bubble ol { margin:.35rem 0 .35rem 1.25rem; }
.chat-bubble table { border-collapse:collapse; width:100%; margin:.35rem 0; }
.chat-bubble table th, .chat-bubble table td { border:1px solid #e5e7eb; padding:.35rem .5rem; }
.chat-bubble a { color:#2563eb; text-decoration:none; } .chat-bubble a:hover { text-decoration:underline; }
.chat-bubble code { background:#f3f4f6; padding:.05rem .25rem; border-radius:4px; }
.chat-bubble pre { background:#111827; color:#f9fafb; padding:.7rem .9rem; border-radius:10px; overflow:auto; font-size:.9rem; }

/* Right editor panel */
.editor-wrap { border:1px solid var(--border); border-radius:10px; padding:.25rem .5rem; }
.ql-container.ql-snow {min-height:360px; border:none;}
.ql-toolbar.ql-snow {border:none; border-bottom:1px solid var(--border);}

/* Replace expander icon safely: hide built-in icon, draw our own */
[data-testid="stExpanderHeader"] [data-testid="stExpanderToggleIcon"]{
  display:none !important; /* prevent ligature text from ever showing */
}
[data-testid="stExpanderHeader"]::before{
  content:"▸";
  display:inline-block;
  margin-right:.35rem;
  color:#6b7280;
  transition: transform .18s ease;
}
[data-testid="stExpanderHeader"][aria-expanded="true"]::before{
  content:"▾";
}
</style>
""", unsafe_allow_html=True)

# ---------- Config ----------
def _gen_id(n=6): return ''.join(random.choices(string.ascii_uppercase + string.digits, k=n))

APP_PASSCODE      = os.getenv("APP_PASSCODE")      or st.secrets.get("env", {}).get("APP_PASSCODE")
ACADEMIC_PASSCODE = os.getenv("ACADEMIC_PASSCODE") or st.secrets.get("env", {}).get("ACADEMIC_PASSCODE")

SPREADSHEET_KEY    = os.getenv("SPREADSHEET_KEY", "1i9kIMnIJkbpOWsqKtcyuTfz-5BREKPNXqESjtWJiDuQ")
ASSIGNMENT_DEFAULT = os.getenv("ASSIGNMENT_ID", "GENERIC")
SIM_THRESHOLD      = float(os.getenv("SIM_THRESHOLD", "0.85"))
AUTO_SAVE_SECONDS  = int(os.getenv("AUTO_SAVE_SECONDS", "60"))

# ---------- Target schemas ----------
EVENTS_HEADERS = ["timestamp", "user_id", "assignment_id", "turn", "prompt", "response"]
DRAFTS_HEADERS = ["user_id", "assignment_id", "draft_html", "draft_text", "last_updated"]

# ---------- Session defaults ----------
st.session_state.setdefault("__auth_ok", False)
st.session_state.setdefault("is_academic", False)
st.session_state.setdefault("user_id", None)
st.session_state.setdefault("show_landing_page", True)
st.session_state.setdefault("assignment_id", ASSIGNMENT_DEFAULT)
st.session_state.setdefault("chat", [])
st.session_state.setdefault("llm_outputs", [])
st.session_state.setdefault("draft_html", "")
st.session_state.setdefault("report", None)
st.session_state.setdefault("last_saved_at", None)
st.session_state.setdefault("last_autosave_at", None)
st.session_state.setdefault("last_saved_html", "")
st.session_state.setdefault("pending_prompt", None)

# ---------- Helpers ----------
def excerpt(text, n=300):
    t = text or ""
    return t if len(t) <= n else t[:n] + " …"

def segment_paragraphs(text: str):
    if not text: return []
    return [p.strip() for p in text.split("\n") if p.strip()]

# Robust Quill wrapper
def render_quill_html(key: str, initial_html: str) -> str:
    try:
        out = st_quill(value=initial_html, placeholder="Write your draft here…", html=True, key=key)
        if isinstance(out, dict) and out.get("html"): return out["html"]
        if isinstance(out, str) and out: return out
    except TypeError:
        try:
            out = st_quill(value=initial_html, placeholder="Write your draft here…", key=key)
        except TypeError:
            out = st_quill(initial_html)
    if isinstance(out, dict):
        if out.get("html"): return out["html"]
        delta = out.get("delta") or out.get("ops") or {}
        ops = delta.get("ops") if isinstance(delta, dict) else delta
        try:
            text = "".join(op.get("insert", "") for op in ops) if isinstance(ops, list) else ""
        except Exception:
            text = ""
        return "<p>" + text.replace("\n", "</p><p>") + "</p>" if text else (initial_html or "")
    if isinstance(out, str): return out
    return initial_html or ""

# ---------- Secrets / clients ----------
def _as_plain(obj):
    """Convert Streamlit AttrDict / nested mappings to plain dicts/lists."""
    if isinstance(obj, Mapping):
        return {k: _as_plain(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [_as_plain(v) for v in obj]
    return obj

@st.cache_resource
def get_gspread_client():
    sa_raw = st.secrets.get("gcp_service_account", None)
    if sa_raw is None:
        sa_raw = os.getenv("GCP_SERVICE_ACCOUNT_JSON")

    if sa_raw is None:
        st.error("GCP Service Account credentials not found in secrets or env (GCP_SERVICE_ACCOUNT_JSON).")
        st.stop()

    if isinstance(sa_raw, str):
        try:
            sa_info = json.loads(sa_raw)
        except json.JSONDecodeError:
            st.error("GCP_SERVICE_ACCOUNT_JSON must be a valid JSON string.")
            st.stop()
    else:
        sa_info = _as_plain(sa_raw)  # handles AttrDict

    scopes = [
        "https://www.googleapis.com/auth/spreadsheets",
        "https://www.googleapis.com/auth/drive",
    ]
    creds = Credentials.from_service_account_info(sa_info, scopes=scopes)
    gc = gspread.authorize(creds)
    return gc.open_by_key(SPREADSHEET_KEY)

@st.cache_resource
def get_llm_client():
    gemini_key = os.getenv("GEMINI_API_KEY") or st.secrets.get("google_api", {}).get("gemini_api_key")
    if not gemini_key:
        st.error("Gemini API key not found in secrets/env.")
        st.stop()
    genai.configure(api_key=gemini_key)
    return genai.GenerativeModel("gemini-1.5-flash")

if genai is None or gspread is None or Credentials is None:
    st.error("Required libraries not available. Please ensure google-generativeai and gspread are installed.")
    st.stop()

sh  = get_gspread_client()
LLM = get_llm_client()

@st.cache_resource
def _get_or_create_ws(title, headers):
    try:
        return sh.worksheet(title)
    except WorksheetNotFound:
        ws = sh.add_worksheet(title=title, rows=1, cols=len(headers))
        ws.append_row(headers, value_input_option="USER_ENTERED")
        return ws

EVENTS_WS = _get_or_create_ws("events", EVENTS_HEADERS)
DRAFTS_WS = _get_or_create_ws("drafts", DRAFTS_HEADERS)

def append_row_safe(ws, row):
    """Robust append that avoids 'Unable to parse range' errors."""
    try:
        ws.append_row(row, value_input_option="USER_ENTERED")
        return
    except Exception:
        pass
    try:
        # Slow but reliable: find next row and write there.
        current = ws.get_all_values()
        next_row = len(current) + 1
        if next_row > ws.row_count:
            ws.add_rows(max(10, next_row - ws.row_count))
        ws.update(f"A{next_row}", [row], value_input_option="USER_ENTERED")
    except Exception as e2:
        st.warning(f"Append failed: {e2}")

# ---------- LLM & logging ----------
def ask_llm(prompt_text: str):
    chunks = []
    try:
        for ch in LLM.generate_content([prompt_text], stream=True):
            if getattr(ch, "text", None):
                chunks.append(ch.text)
    except Exception as e:
        chunks.append(f"Error: {e}")
    return "".join(chunks)

def log_turn(prompt: str, response: str):
    """One row per turn (timestamp, user, assignment, turn, prompt, response)."""
    turn = sum(1 for m in st.session_state.chat if m["role"] == "user")
    append_row_safe(EVENTS_WS, [
        datetime.datetime.now().isoformat(),
        st.session_state.user_id,
        st.session_state.assignment_id,
        turn,
        excerpt(prompt, 10000),
        excerpt(response, 10000),
    ])

def log_ping(note="login"):
    append_row_safe(EVENTS_WS, [
        datetime.datetime.now().isoformat(),
        st.session_state.user_id or "",
        st.session_state.assignment_id or "",
        0,
        note,
        "",
    ])

def save_progress(silent=False):
    draft_text = html_to_text(st.session_state.draft_html)
    append_row_safe(DRAFTS_WS, [
        st.session_state.user_id,
        st.session_state.assignment_id,
        st.session_state.draft_html,
        draft_text,
        datetime.datetime.now().isoformat()
    ])
    st.session_state.last_saved_at = datetime.datetime.now()
    st.session_state.last_saved_html = st.session_state.draft_html
    if not silent: st.toast("Draft saved")

def load_progress():
    try:
        recs = DRAFTS_WS.get_all_records(expected_headers=DRAFTS_HEADERS, head=1, default_blank="")
        for r in reversed(recs):
            if str(r.get("user_id","")).strip().upper() == st.session_state.user_id.strip().upper() and \
               str(r.get("assignment_id","")).strip() == st.session_state.assignment_id.strip():
                return r.get("draft_html") or ""
    except Exception:
        return ""
    return ""

def maybe_autosave():
    now = time.time()
    last_ts = st.session_state.last_autosave_at or 0
    changed = (st.session_state.draft_html or "") != (st.session_state.last_saved_html or "")
    if changed and (now - last_ts) >= AUTO_SAVE_SECONDS:
        save_progress(silent=True)
        st.session_state.last_autosave_at = now

# ---------- Similarity ----------
def compute_similarity_report(final_text, llm_texts, sim_thresh=SIM_THRESHOLD):
    finals = segment_paragraphs(final_text)
    llm_segs = [s for t in llm_texts for s in segment_paragraphs(t)]
    if not finals or not llm_segs:
        return {"backend": SIM_BACKEND, "mean": 0.0, "high_share": 0.0, "rows": []}

    rows, high_tokens = [], 0
    total_tokens = sum(len(s.split()) for s in finals)

    if SIM_BACKEND == "sbert" and _sbert_model is not None:
        Ef = _sbert_model.encode(finals, convert_to_tensor=True, normalize_embeddings=True)
        El = _sbert_model.encode(llm_segs, convert_to_tensor=True, normalize_embeddings=True)
        sims = sbert_util.cos_sim(Ef, El).cpu().numpy()
        for i, fseg in enumerate(finals):
            j = int(sims[i].argmax()); s = float(sims[i, j]); nearest = llm_segs[j]
            rows.append({"final_seg": excerpt(fseg, 200), "nearest_llm": excerpt(nearest, 200), "cosine": round(s, 3)})
            if s >= sim_thresh: high_tokens += len(fseg.split())
    elif SIM_BACKEND == "tfidf":
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        vectorizer = TfidfVectorizer().fit(finals + llm_segs)
        F = vectorizer.transform(finals); L = vectorizer.transform(llm_segs)
        sims = cosine_similarity(F, L)
        for i, fseg in enumerate(finals):
            j = int(sims[i].argmax()); s = float(sims[i, j]); nearest = llm_segs[j]
            rows.append({"final_seg": excerpt(fseg, 200), "nearest_llm": excerpt(nearest, 200), "cosine": round(s, 3)})
            if s >= sim_thresh: high_tokens += len(fseg.split())
    else:
        from difflib import SequenceMatcher
        def cos_like(a, b): return SequenceMatcher(None, a, b).ratio()
        for fseg in finals:
            best, nearest = 0.0, ""
            for l in llm_segs:
                c = cos_like(fseg, l)
                if c > best: best, nearest = c, l
            rows.append({"final_seg": excerpt(fseg, 200), "nearest_llm": excerpt(nearest, 200), "cosine": round(best, 3)})
            if best >= sim_thresh: high_tokens += len(fseg.split())

    mean_sim = 0.0 if not rows else round(sum(r["cosine"] for r in rows) / len(rows), 3)
    high_share = round(high_tokens / max(1, total_tokens), 3)
    return {"backend": SIM_BACKEND, "mean": mean_sim, "high_share": high_share, "rows": rows[:30]}

# ---------- Evidence export ----------
def export_evidence_docx(user_id, assignment_id, chat, draft_html, report):
    if not DOCX_OK:
        raise RuntimeError("python-docx not installed")
    final_text = html_to_text(draft_html)
    d = docx.Document()
    d.add_heading("Coursework Evidence Pack", 0)
    d.add_paragraph(f"User ID: {user_id}")
    d.add_paragraph(f"Assignment ID: {assignment_id}")
    d.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    d.add_heading("Chat with LLM", level=1)
    for m in chat:
        role = "Student" if m["role"] == "user" else "LLM"
        d.add_paragraph(f"{role}: {m['text']}")
    d.add_heading("Final Draft (plain text extract)", level=1)
    for para in final_text.split("\n"):
        d.add_paragraph(para)
    d.add_heading("Similarity Report", level=1)
    d.add_paragraph(f"Backend: {report.get('backend','-')}")
    d.add_paragraph(f"Mean similarity: {report.get('mean',0.0)}")
    d.add_paragraph(f"High-sim share: {report.get('high_share',0.0)*100:.1f}%")
    for r in report.get("rows", []):
        d.add_paragraph(f"- Cosine: {r['cosine']} | Final: {r['final_seg']} | LLM: {r['nearest_llm']}")
    buf = io.BytesIO(); d.save(buf); buf.seek(0); return buf.read()

# ---------- Login / Landing ----------
def login_view():
    # Big hero
    st.markdown('<div class="hero">', unsafe_allow_html=True)
    st.markdown("<h1>LLM Coursework Helper</h1>", unsafe_allow_html=True)
    st.markdown(
        "<p>This pilot helps students ideate & write with an AI assistant while giving academics a transparent view of "
        "<strong>process evidence</strong> (chat turns & draft evolution).</p>",
        unsafe_allow_html=True
    )
    st.markdown('</div>', unsafe_allow_html=True)

    # Two clean cards
    st.markdown('<div class="info-grid">', unsafe_allow_html=True)

    # Students card
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown('<h3>For Students <span class="badge">pilot</span></h3>', unsafe_allow_html=True)
    st.markdown("""
- Brainstorm with the AI (all turns are logged).
- Draft in the rich editor; autosave & resume anytime.
- Run **similarity check** vs AI outputs to keep your own voice.
- Export a **DOCX evidence pack** (chat + draft + similarity).
- **Responsibility:** follow your assessment rules and cite sources.
    """)
    st.markdown('</div>', unsafe_allow_html=True)

    # Academics card
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown('<h3>For Academics</h3>', unsafe_allow_html=True)
    st.markdown("""
- Dashboard to review students’ **turn-by-turn** interactions.
- See **latest draft** and AI ↔ student exchanges.
- Optional writing-alignment indicator for oversight (not grading).
- Data stored: timestamp, pseudonymous student ID, assignment ID, prompt, response, latest draft snapshot.
- Data minimisation: no personal identifiers beyond the provided ID.
    """)
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)  # close grid

    with st.expander("Privacy & Ethics"):
        st.markdown("""
We log **timestamps**, your **Student ID**, **Assignment ID**, each **prompt**, the **AI response**, and saved **draft snapshots**.
We do **not** log personal identifiers beyond the ID you enter. The AI output is advisory; you are responsible for citation and originality.
        """)

    st.markdown("---")

    @st.cache_data(ttl=60)
    def known_student_ids():
        try:
            recs = DRAFTS_WS.get_all_records(expected_headers=DRAFTS_HEADERS, head=1, default_blank="")
            if not recs: return set()
            return set(pd.DataFrame(recs)["user_id"].astype(str).unique())
        except Exception:
            return set()

    user_input = st.text_input(
        "Enter your **Student ID** or a **Passcode**",
        placeholder="Student ID, Student Passcode, or Academic Passcode",
    )

    c1, c2 = st.columns([1, 1])
    with c1:
        if st.button("Login", use_container_width=True):
            inp = (user_input or "").strip().upper()
            if inp and ACADEMIC_PASSCODE and inp == ACADEMIC_PASSCODE.upper():
                st.session_state.update({"__auth_ok": True, "is_academic": True, "user_id": "Academic", "show_landing_page": False})
                log_ping("academic_login")
                st.rerun()
            elif inp and APP_PASSCODE and inp == APP_PASSCODE.upper():
                new_id = _gen_id()
                st.session_state.update({"__auth_ok": True, "is_academic": False, "user_id": new_id, "show_landing_page": True})
                log_ping("student_newid_login")
                st.success(f"Your new Student ID is **{new_id}** — copy it to resume later.")
                st.rerun()
            elif inp in known_student_ids():
                st.session_state.update({"__auth_ok": True, "is_academic": False, "user_id": inp, "show_landing_page": False})
                log_ping("student_return_login")
                st.rerun()
            else:
                st.error("Invalid ID or Passcode.")
    with c2:
        if st.button("Generate New Student ID", use_container_width=True):
            new_id = _gen_id()
            st.session_state.update({"__auth_ok": True, "is_academic": False, "user_id": new_id, "show_landing_page": True})
            log_ping("student_newid_login")
            st.success(f"Your new Student ID is **{new_id}** — copy it to resume later.")
            st.rerun()

# ---------- Academic dashboard ----------
def render_academic_dashboard():
    st.title("🎓 Academic Dashboard")

    @st.cache_data(ttl=300)
    def get_all_student_data():
        try:
            drafts = pd.DataFrame(DRAFTS_WS.get_all_records(
                expected_headers=DRAFTS_HEADERS, head=1, default_blank=""
            ))
            events = pd.DataFrame(EVENTS_WS.get_all_records(
                expected_headers=EVENTS_HEADERS, head=1, default_blank=""
            ))
            return drafts, events
        except Exception as e:
            st.error(f"Could not fetch data from Google Sheets: {e}")
            return pd.DataFrame(), pd.DataFrame()

    drafts_df, events_df = get_all_student_data()
    if drafts_df.empty and events_df.empty:
        st.warning("No student data recorded yet.")
        return

    all_ids = sorted({*(drafts_df.get('user_id', pd.Series([])).dropna().astype(str)),
                      *(events_df.get('user_id', pd.Series([])).dropna().astype(str))} - {"Academic"})

    sid = st.selectbox("Select a Student ID to review", all_ids, index=None, placeholder="Search…")
    if not sid:
        st.info("Select a student to begin.")
        return

    st.header(f"Reviewing: {sid}")
    s_drafts = drafts_df[drafts_df['user_id'] == sid]
    s_events = events_df[events_df['user_id'] == sid].sort_values("timestamp")

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Latest Draft")
        if not s_drafts.empty:
            latest = s_drafts.sort_values('last_updated', ascending=False).iloc[0]
            st.markdown(f"**Last Saved:** {latest['last_updated']}")
            st_html(f'<div class="chat-box">{latest["draft_html"]}</div>', height=600)
        else:
            st.info("No saved drafts for this student.")

    with col2:
        st.subheader("Turns (Prompt → Response)")
        if not s_events.empty:
            bubbles = []
            for _, row in s_events.iterrows():
                bubbles.append(f'<div class="chat-bubble chat-user"><strong>Prompt (Turn {row.get("turn","?")}):</strong><br>{md_to_html(row["prompt"])}</div>')
                bubbles.append(f'<div class="chat-bubble chat-assistant"><strong>Response:</strong><br>{md_to_html(row["response"])}</div>')
            st_html(f'<div class="chat-box">{"".join(bubbles)}</div>', height=600)
        else:
            st.info("No chat history for this student.")

# ---------- Student view ----------
def render_student_view():
    if st.session_state.get("show_landing_page", False):
        st.markdown('<div class="hero">', unsafe_allow_html=True)
        st.markdown("<h1>Welcome</h1>", unsafe_allow_html=True)
        st.markdown("Use the AI to brainstorm and the editor to write. All turns & draft saves are logged for **process evidence**.")
        st.markdown('</div>', unsafe_allow_html=True)
        with st.expander("What gets logged?"):
            st.markdown("- **Timestamps**, your **Student ID**, **Assignment ID**, each **prompt**, the **AI response**, and saved **draft snapshots**.")
        st.markdown("---")
        c1, c2 = st.columns([1,1])
        with c1:
            if st.button("Get started", type="primary", use_container_width=True):
                st.session_state.show_landing_page = False
                st.rerun()
        with c2:
            if st.button("Sign out", use_container_width=True):
                for k in list(st.session_state.keys()):
                    del st.session_state[k]
                st.rerun()
        return

    # Header
    st.markdown(
        f'<div class="header-bar">'
        f'<div class="status-chip">User: {st.session_state.user_id}</div>'
        f'<div class="status-chip">Assignment: {st.session_state.assignment_id}</div>'
        f'<div class="status-chip">Similarity backend: {SIM_BACKEND}</div>'
        f'<div class="small-muted">Last saved: {st.session_state.last_saved_at.strftime("%H:%M:%S") if st.session_state.last_saved_at else "—"}</div>'
        f'</div>', unsafe_allow_html=True
    )

    # Toolbar
    t1, t2, t3 = st.columns([1.2, 0.9, 0.8])
    with t1:
        st.session_state.assignment_id = st.text_input("Assignment ID", value=st.session_state.assignment_id)
    with t2:
        if st.button("🔄 Load Last Draft"):
            html = load_progress()
            if html:
                st.session_state.draft_html = html
                st.success("Loaded last saved draft.")
                st.rerun()
            else:
                st.warning("No saved draft found.")
    with t3:
        if st.button("🧹 Clear Chat"):
            st.session_state.chat = []; st.session_state.llm_outputs = []; st.toast("Chat cleared")

    left, right = st.columns([0.5, 0.5], gap="large")

    # Left: Assistant
    with left:
        st.subheader("💬 Assistant")

        # Build bubbles from history
        if not st.session_state.get("chat"):
            bubbles_html = '<div class="chat-empty">Ask for ideas, critique, or examples.</div>'
        else:
            all_bubbles = []
            for m in st.session_state.chat:
                css = "chat-user" if m["role"] == "user" else "chat-assistant"
                content = md_to_html(m.get("text", "")) if m["role"] == "assistant" \
                          else _html.escape(m.get("text", "")).replace("\n", "<br>")
                all_bubbles.append(f'<div class="chat-bubble {css}">{content}</div>')
            bubbles_html = "".join(all_bubbles)

        if st.session_state.get("pending_prompt"):
            bubbles_html += '<div class="chat-bubble chat-assistant">…thinking</div>'

        st_html(
            f'<div id="chatbox" class="chat-box">{bubbles_html}</div>'
            f'<script>var b=document.getElementById("chatbox"); if(b) b.scrollTop=b.scrollHeight;</script>',
            height=600
        )

        with st.form("chat_form", clear_on_submit=True):
            c1, c2 = st.columns([4, 1])
            with c1:
                prompt = st.text_input("Ask…", "", placeholder="Type and press Send", label_visibility="collapsed")
            with c2:
                send = st.form_submit_button("Send")

        if send and (prompt or "").strip():
            st.session_state.chat.append({"role": "user", "text": prompt})
            st.session_state.pending_prompt = prompt
            st.rerun()

        if st.session_state.get("pending_prompt"):
            with st.spinner("Generating response…"):
                p = st.session_state.pending_prompt
                st.session_state.pending_prompt = None
                reply = ask_llm(p)
                st.session_state.chat.append({"role": "assistant", "text": reply})
                st.session_state.llm_outputs.append(reply)
                log_turn(prompt=p, response=reply)
            st.rerun()

    # Right: Draft (KPIs removed for speed)
    with right:
        st.subheader("📝 Draft")
        st.markdown('<div class="editor-wrap">', unsafe_allow_html=True)
        st.session_state.draft_html = render_quill_html("editor", st.session_state.draft_html)
        st.markdown('</div>', unsafe_allow_html=True)

        maybe_autosave()

        c1, c2, c3 = st.columns(3)
        with c1:
            if st.button("💾 Save Draft", use_container_width=True):
                save_progress()
        with c2:
            if st.button("📊 Run Similarity", use_container_width=True):
                plain = html_to_text(st.session_state.draft_html)
                if plain.strip() and st.session_state.llm_outputs:
                    report = compute_similarity_report(plain, st.session_state.llm_outputs, SIM_THRESHOLD)
                    st.session_state.report = report
                    st.success(f"Mean: {report['mean']} | High-sim: {report['high_share']*100:.1f}%")
                    with st.expander("Matches (trimmed)"):
                        for r in report["rows"]:
                            st.markdown(f"- **Cos:** {r['cosine']}  \n  **Final:** {r['final_seg']}  \n  **LLM:** {r['nearest_llm']}")
                else:
                    st.warning("Need draft text + at least one LLM response.")
        with c3:
            if st.button("⬇️ Export Evidence (DOCX)", use_container_width=True):
                try:
                    rep = st.session_state.get("report", {"backend": SIM_BACKEND, "mean": 0.0, "high_share": 0.0, "rows": []})
                    data = export_evidence_docx(
                        st.session_state.user_id,
                        st.session_state.assignment_id,
                        st.session_state.chat,              # correct chat passed
                        st.session_state.draft_html,
                        rep
                    )
                    st.download_button(
                        "Download DOCX",
                        data=data,
                        file_name=f"evidence_{st.session_state.user_id}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"Export failed: {e}")

# ---------- Router ----------
if not st.session_state["__auth_ok"]:
    login_view()
else:
    if st.session_state.get("is_academic"):
        render_academic_dashboard()
    else:
        render_student_view()
